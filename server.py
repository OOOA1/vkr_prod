# server.py
# Зависимости:
#   pip install fastapi "uvicorn[standard]" python-multipart pandas openpyxl docxtpl requests
# (docxtpl тянет python-docx, используется для генерации DOCX-инструкции)

import io
import re
import csv
import zipfile
from pathlib import Path
from typing import Optional, Dict, Tuple, List

import os
import tempfile
import subprocess

import pandas as pd
import requests
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Query
from fastapi.responses import (
    HTMLResponse,
    JSONResponse,
    StreamingResponse,
    PlainTextResponse,
    FileResponse,
)
from docxtpl import DocxTemplate
import jinja2
JINJA_ENV = jinja2.Environment()

# для генерации DOCX-инструкции (ставится вместе с docxtpl)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment

from templates_config import TEMPLATES
import unicodedata



app = FastAPI(title="Help University — DOCX → ZIP", version="3.6.0")

# === Стабильные ID для шаблонов ===
def slug_id(v: str) -> str:
    v = unicodedata.normalize("NFKC", v)
    v = re.sub(r"\s+", " ", v).strip()
    v = v.replace("\\", "/")
    allowed = "._-() "
    out = []
    for ch in v:
        out.append(ch if (ch.isalnum() or ch in allowed) else "_")
    return re.sub(r"\s+", "_", "".join(out)).lower()

# один раз навешиваем id на все шаблоны (учитываем ПУТЬ, чтобы комплекты не пересекались)
for idx, tpl in enumerate(TEMPLATES):
    if "id" not in tpl:
        # "input/first/дневник.docx" -> "input/first/дневник"
        rel = tpl["path"].replace("\\", "/")
        rel_no_ext = re.sub(r"\.[^.\\/]+$", "", rel)
        tpl["id"] = slug_id(rel_no_ext) or f"tpl_{idx:03d}"

@app.get("/catalog")
def catalog(prefix: Optional[str] = None):
    """
    Отдаём список документов. Если передан prefix,
    фильтруем только шаблоны, у которых path начинается с этого префикса.
    """
    tpls = TEMPLATES
    if prefix:
        pfx = prefix.replace("\\", "/")
        tpls = [
            t for t in TEMPLATES
            if t["path"].replace("\\", "/").startswith(pfx)
        ]

    items = []
    for t in tpls:
        path_norm = t["path"].replace("\\", "/")
        items.append({
            "id": t["id"],
            "title": Path(path_norm).stem + ".docx",
            "path": path_norm,
        })
    return {"items": items}

# === Пути базы ===
BASE_DIR = Path(__file__).resolve().parent

# === Настройки выдачи Excel-шаблона ===
TEMPLATE_DOWNLOAD_NAME = "main_example.xlsx"
CANDIDATE_TEMPLATE_PATHS: List[Path] = [
    BASE_DIR / "main_example.xlsx",
    BASE_DIR / "main.xlsx",
    BASE_DIR / "main — копия.xlsx",
]

# === Готовые Excel-шаблоны под каждый комплект ===
# Папку "table_templates" создай рядом с server.py и положи туда свои 4 файла.
# Ключи (kit1, kit2, ...) ДОЛЖНЫ совпадать со значениями <option value="..."> в <select id="direction">.
KIT_TEMPLATES: Dict[str, Path] = {
    "kit1": BASE_DIR / "table_templates" / "First шаблон.xlsx",
    "kit2": BASE_DIR / "table_templates" / "Менеджмент УП экономика шаблон.xlsx",
    "kit3": BASE_DIR / "table_templates" / "Реклама, лингвистика, журналистика, ГМУ шаблон.xlsx",
    "kit4": BASE_DIR / "table_templates" / "docx11 шаблон.xlsx",
}

# === Настройки выдачи Инструкции ===
INSTRUCTION_DOWNLOAD_NAME = "instruction.docx"
INSTRUCTION_CANDIDATES: List[Path] = [
    BASE_DIR / "instruction.docx",
    BASE_DIR / "инструкция.docx",
    BASE_DIR / "instruction.doc",   # если вдруг положите .doc
    BASE_DIR / "instruction.dock",  # опечатка — тоже поддержим
]

# ============= Красивый UI (без внешних зависимостей) =============
INDEX_HTML = """
<!doctype html>
<html lang="ru">
<head>
<meta charset="utf-8">
<title>Help University — Автоматизация документов</title>
<meta name="viewport" content="width=device-width, initial-scale=1">

<style>
  :root {
    --brand: #3b82f6;
    --brand-glow: #60a5fa;
    --bg: #0f172a;
    --card: rgba(17,25,40,0.85);
    --text: #f8fafc;
    --subtext: #94a3b8;
    --radius: 18px;
    --blur: 20px;
  }

  * { box-sizing: border-box; }

  body {
    margin: 0;
    font-family: "Inter", system-ui, sans-serif;
    color: var(--text);
    background: radial-gradient(circle at 30% 10%, #1e3a8a 0%, #0f172a 80%);
    overflow-x: hidden;
    animation: fadeInBg 2s ease;
  }

  @keyframes fadeInBg {
    from {opacity: 0;}
    to {opacity: 1;}
  }

  header {
    position: sticky;
    top: 0;
    z-index: 10;
    backdrop-filter: blur(var(--blur));
    background: rgba(17,25,40,0.65);
    border-bottom: 1px solid rgba(255,255,255,0.05);
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 20px 40px;
  }

  .logo {
    font-weight: 800;
    font-size: 26px;
    background: linear-gradient(90deg, var(--brand-glow), #7dd3fc);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    letter-spacing: -0.5px;
  }

  .tagline {
    font-size: 14px;
    color: var(--subtext);
  }

  main {
    max-width: 920px;
    margin: 70px auto;
    padding: 40px 50px;
    border-radius: var(--radius);
    background: var(--card);
    backdrop-filter: blur(var(--blur));
    box-shadow: 0 0 60px rgba(59,130,246,0.15);
    animation: slideUp 0.8s ease;
  }

  @keyframes slideUp {
    from {opacity:0; transform: translateY(20px);}
    to {opacity:1; transform: translateY(0);}
  }

  h1 {
    font-size: 32px;
    margin-bottom: 10px;
    background: linear-gradient(90deg, var(--brand-glow), #93c5fd);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
  }

  p.sub {
    color: var(--subtext);
    font-size: 15px;
    margin-top: 0;
  }

  label {
    font-weight: 600;
    display: block;
    margin-top: 24px;
  }

    select, input[type=file], input[type=url] {
        width: 100%;
        margin-top: 8px;
        padding: 14px;
        border-radius: 12px;
        border: 1px solid rgba(255,255,255,0.1);
        background: rgba(255,255,255,0.05);
        color: var(--text);
        font-size: 15px;
        transition: all 0.3s ease;
    }

    select option {
        color: #000;
        background: #fff;
    }

  select:focus, input:focus {
    outline: none;
    border-color: var(--brand);
    box-shadow: 0 0 10px rgba(59,130,246,0.3);
    background: rgba(255,255,255,0.08);
  }

  .row {
    display: flex;
    flex-wrap: wrap;
    gap: 12px;
    margin-top: 20px;
  }

  button {
    flex: 1;
    padding: 14px 18px;
    font-size: 15px;
    border-radius: 12px;
    border: none;
    cursor: pointer;
    font-weight: 600;
    transition: all 0.3s ease;
  }

  .btn-primary {
    background: linear-gradient(90deg, var(--brand), var(--brand-glow));
    color: white;
    box-shadow: 0 0 20px rgba(59,130,246,0.25);
  }

  .btn-primary:hover {
    transform: translateY(-2px);
    box-shadow: 0 0 30px rgba(96,165,250,0.4);
  }

  .btn-outline {
    background: transparent;
    border: 2px solid var(--brand);
    color: var(--brand-glow);
  }

  .btn-outline:hover {
    background: var(--brand);
    color: white;
    box-shadow: 0 0 25px rgba(59,130,246,0.4);
  }

  .divider {
    height: 1px;
    background: rgba(255,255,255,0.1);
    margin: 36px 0;
  }

  .docs {
    margin-top: 18px;
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(260px, 1fr));
    gap: 12px;
  }

  .doc-item {
    background: rgba(255,255,255,0.05);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 12px;
    padding: 10px 14px;
    display: flex;
    align-items: center;
    gap: 10px;
    transition: background .3s, transform .2s, box-shadow .2s;
  }

  .doc-item:hover {
    background: rgba(255,255,255,0.08);
    transform: translateY(-2px);
    box-shadow: 0 0 12px rgba(59,130,246,0.2);
  }

  footer {
    text-align: center;
    margin: 60px 0 20px;
    color: var(--subtext);
    font-size: 14px;
  }

  .glow {
    position: absolute;
    inset: 0;
    pointer-events: none;
    background: radial-gradient(circle at 30% 0%, rgba(59,130,246,0.15), transparent 70%);
    z-index: -1;
  }

  /* общий стиль поля select */
select {
  width: 100%;
  margin-top: 8px;
  padding: 14px 16px;
  border-radius: 12px;
  border: 1px solid rgba(255,255,255,0.1);
  background: linear-gradient(145deg, rgba(30,41,59,0.9), rgba(17,25,40,0.9));
  color: var(--text);
  font-size: 15px;
  appearance: none; /* убираем стандартную стрелку */
  -webkit-appearance: none;
  -moz-appearance: none;
  background-image: url('data:image/svg+xml;utf8,<svg fill="white" height="16" width="16" xmlns="http://www.w3.org/2000/svg"><path d="M4 6l4 4 4-4z"/></svg>');
  background-repeat: no-repeat;
  background-position: right 12px center;
  background-size: 14px;
  transition: all 0.3s ease;
}

select:hover {
  border-color: rgba(96,165,250,0.5);
  background: linear-gradient(145deg, rgba(37,54,84,0.95), rgba(20,29,50,0.95));
}

select:focus {
  outline: none;
  border-color: var(--brand);
  box-shadow: 0 0 10px rgba(59,130,246,0.4);
  background: linear-gradient(145deg, rgba(40,60,90,0.95), rgba(22,32,55,0.95));
}

/* оформление выпадающих опций */
select option {
  background: #1e293b;
  color: #f8fafc;
  padding: 10px;
  border: none;
}

/* подсветка при наведении на вариант */
select option:hover {
  background: #2563eb;
  color: white;
}


  @media (max-width:600px){
    main {padding: 25px;}
  }

  /* subtle floating animation for buttons */
  .floaty {
    animation: float 3s ease-in-out infinite;
  }
  @keyframes float {
    0%,100% {transform: translateY(0);}
    50% {transform: translateY(-4px);}
  }
</style>
</head>
<body>

<div class="glow"></div>

<header>
  <div class="logo">Help University</div>
  <div class="tagline">Автоматизация документов</div>
</header>

<main>
  <h1>Автоматизация документов</h1>
  <p class="sub">Создавайте и скачивайте шаблоны по направлениям подготовки. Лёгкий, современный и стильный интерфейс.</p>

  <label>Загрузить таблицу Excel или CSV</label>
  <input type="file" id="fileInput" accept=".xlsx,.csv">
  <small style="color:var(--subtext)">Поддерживаются .xlsx и .csv файлы</small>

  <label style="margin-top:14px;">Или вставьте ссылку на Google Sheet</label>
  <input type="url" id="gsheetUrl" placeholder="https://docs.google.com/spreadsheets/d/...">

  <div class="row">
    <button class="btn-outline floaty" id="btnTemplate">📄 Скачать шаблон</button>
    <button class="btn-outline floaty" id="btnInstruction">📘 Скачать инструкцию</button>
  </div>

  <div class="divider"></div>

  <label for="direction">Направление подготовки</label>
  <select id="direction">
    <option value="">— выберите комплект —</option>
    <option value="kit1">Условный комплект 1(first)</option>
    <option value="kit2">Менеджмент УП экономика</option>
    <option value="kit3">Реклама, лингвистика, журналистика, ГМУ</option>
    <option value="kit4">Условный комплект 4(new_docx11)</option>
    <!-- добавишь ещё, когда появятся новые наборы -->
  </select>

  <div id="docs" class="docs"></div>

  <div class="row" style="margin-top:28px;">
    <button class="btn-primary floaty" id="downloadBtn" disabled>⬇️ Сгенерировать ZIP</button>
  </div>
</main>

<footer>© 2025 Help University • Интеллектуальная автоматизация документов</footer>

<script>
  // Комплекты → папка в input/
  // ВАЖНО: здесь должны быть реальные папки, в которых лежат шаблоны из templates_config.py
  const kitFolders = {
    // пример для будущего:
    // kitn: "input/new_docx11111/",
    kit1: "input/first/",
    kit2: "input/менеджмент_УП_экономика",
    kit3: "input/Реклама, лингвистика, журналистика, ГМУ",
    kit4: "input/new_docx11/",
  };

  // НОВОЕ: соответствие "комплект → имя Excel-шаблона"
  // Здесь должны быть ТОЧНО такие же имена, как в KIT_TEMPLATES на бэкенде.
  const kitTemplateNames = {
    kit1: "First шаблон",
    kit2: "Менеджмент УП экономика шаблон",
    kit3: "Реклама, лингвистика, журналистика, ГМУ.xlsx",
    kit4: "docx11 шаблон",
  };

  const directionSelect = document.getElementById("direction");
  const docsDiv = document.getElementById("docs");
  const downloadBtn = document.getElementById("downloadBtn");
  const fileInput = document.getElementById("fileInput");
  const gsheetUrl = document.getElementById("gsheetUrl");

  async function loadKitDocs(kit) {
    docsDiv.innerHTML = "";
    if (!kit) {
      downloadBtn.disabled = true;
      return;
    }

    const prefix = kitFolders[kit];
    if (!prefix) {
      downloadBtn.disabled = true;
      docsDiv.innerHTML = '<div class="empty">Для этого комплекта ещё не настроена папка</div>';
      return;
    }

    downloadBtn.disabled = false;

    try {
      const resp = await fetch("/catalog?prefix=" + encodeURIComponent(prefix));
      if (!resp.ok) throw new Error("HTTP " + resp.status);
      const data = await resp.json();
      const items = data.items || [];

      if (!items.length) {
        docsDiv.innerHTML = '<div class="empty">В этой папке пока нет шаблонов</div>';
        return;
      }

      items.forEach(doc => {
        const item = document.createElement("div");
        item.className = "doc-item";
        item.dataset.id = doc.id;
        item.textContent = doc.title;
        docsDiv.appendChild(item);
      });
    } catch (e) {
      console.error(e);
      docsDiv.innerHTML = '<div class="empty">Не удалось загрузить список документов</div>';
    }
  }

  directionSelect.addEventListener("change", () => {
    const kit = directionSelect.value;
    loadKitDocs(kit);
  });

  function blobDownload(filename, blob){
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url;
    a.download = filename;
    document.body.appendChild(a);
    a.click();
    a.remove();
    URL.revokeObjectURL(url);
  }

   // ==== скачать шаблон ТОЛЬКО по выбранному комплекту (без генерации по полям) ====
  document.getElementById("btnTemplate").addEventListener("click", async () => {
    const kit = directionSelect.value;
    if (!kit) {
      alert("Сначала выберите комплект");
      return;
    }

    // теперь отправляем только id комплекта, а не список документов
    const url = "/template?kit=" + encodeURIComponent(kit);

    try {
      const resp = await fetch(url);
      if (!resp.ok) {
        const text = await resp.text();
        alert("Ошибка при скачивании шаблона: " + text);
        return;
      }

      const blob = await resp.blob();

      // имя файла берём из словаря, чтобы совпадало с реальным шаблоном
      const filename = kitTemplateNames[kit] || "template.xlsx";
      blobDownload(filename, blob);
    } catch (e) {
      alert("Сетевая ошибка при скачивании шаблона: " + e.message);
    }
  });

  // ==== старый функционал: скачать инструкцию ====
  document.getElementById("btnInstruction").addEventListener("click", () => {
    window.location.href = "/instruction";
  });

  // ==== старый функционал: сгенерировать ZIP по таблице ====
   // ==== генерация ZIP: ТОЛЬКО выбранный комплект и выбранные документы ====
  downloadBtn.addEventListener("click", async () => {
    const kit = directionSelect.value;
    if (!kit) {
      alert("Выберите комплект!");
      return;
    }

    // Берём все документы текущего комплекта (без галочек — все пойдут в ZIP)
    const items = [...docsDiv.querySelectorAll('.doc-item')];
    if (!items.length) {
      alert("В этом комплекте нет документов");
      return;
    }
    const ids = items
      .map(el => el.dataset.id)
      .filter(Boolean);

    // Источник данных: файл или Google Sheet
    const hasFile = fileInput.files && fileInput.files[0];
    const gsheet = gsheetUrl.value.trim();

    if (!hasFile && !gsheet) {
      alert("Загрузите файл или вставьте ссылку на Google Sheet");
      return;
    }

    const fd = new FormData();
    if (hasFile) {
      fd.append("table_file", fileInput.files[0]);      // как и раньше
    } else {
      fd.append("gsheet_url", gsheet);                 // как и раньше
    }
    fd.append("header_row", "1");                       // как в старом UI
    fd.append("include", ids.join(","));                // КЛЮЧЕВОЕ: список id шаблонов

    const prevText = downloadBtn.textContent;
    downloadBtn.disabled = true;
    downloadBtn.textContent = "⏳ Генерация...";

    try {
      const resp = await fetch("/generate", {
        method: "POST",
        body: fd
      });
      if (!resp.ok) {
        let msg = `HTTP ${resp.status}`;
        try {
          const data = await resp.json();
          msg = data.detail || data.error || msg;
        } catch (_) {}
        throw new Error(msg);
      }
      const blob = await resp.blob();
      blobDownload("generated_docs.zip", blob);
    } catch (e) {
      alert("Ошибка при генерации: " + e.message);
    } finally {
      downloadBtn.disabled = false;
      downloadBtn.textContent = prevText;
    }
  });
</script>

</body>
</html>
"""

# ============= Бизнес-логика =============
INVALID_FS = r'[<>:"/\\|?*]'

def safe(v): return "" if (v is None or pd.isna(v)) else str(v).strip()

def letter(value: str, index: int) -> str:
    """
    Берём строку (ФИО), убираем пробелы и возвращаем букву по индексу.
    Если букв меньше либо индекс вне диапазона — вернём пустую строку.
    """
    s = safe(value or "")
    # убираем пробелы и неразрывные пробелы
    s = re.sub(r"\s+", "", s).replace("\xa0", "")
    if not s:
        return ""
    if 0 <= index < len(s):
        return s[index].upper()
    return ""

def lc(value: str) -> str:
    """
    Принудительно переводим строку в нижний регистр.
    Используем в шаблоне как {{ Поле|lc }}.
    """
    return safe(value).lower()

def uc(value: str) -> str:
    """
    Переводит строку в ВЕРХНИЙ регистр.
    Использование в шаблоне: {{ Поле|uc }}
    """
    return safe(value).upper()

JINJA_ENV.filters["letter"] = letter
JINJA_ENV.filters["lc"] = lc
JINJA_ENV.filters["uc"] = uc

class SafeMap(dict):
    def __missing__(self, key): return ""

def slugify(name: str) -> str:
    return re.sub(INVALID_FS, "_", name).rstrip(" .") or "file"

def slugify_path(path: str) -> str:
    parts = re.split(r"[\\/]+", (path or "").strip())
    parts = [slugify(p) for p in parts if p and p.strip()]
    return "/".join(parts)

SOFFICE_BIN = os.getenv("SOFFICE_BIN", "soffice")  # на Windows можно указать полный путь до soffice.exe

def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Конвертирует DOCX (bytes) -> PDF (bytes) через LibreOffice (soffice --headless).
    Используем тот же стиль, как ты запускал вручную из консоли.
    """
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)

        in_path = td / "input.docx"
        out_dir = td / "out"

        out_dir.mkdir(parents=True, exist_ok=True)
        in_path.write_bytes(docx_bytes)

        cmd = [
            SOFFICE_BIN,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(out_dir),
            str(in_path),
        ]

        proc = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )

        # Ищем любой PDF, который LibreOffice сгенерировал
        pdf_files = list(out_dir.glob("*.pdf"))

        if proc.returncode != 0 or not pdf_files:
            raise RuntimeError("LibreOffice DOCX→PDF failed:\n" + (proc.stdout or ""))

        # Берём первый найденный PDF
        return pdf_files[0].read_bytes()

def _norm(s: str) -> str:
    return re.sub(r"\s+", "", str(s)).replace("\ufeff","").replace("\xa0","").replace("ё","е").lower()

def expected_headers() -> set:
    exp = {"фио","группа"}
    for tpl in TEMPLATES:
        exp |= {_norm(v) for v in tpl["fields"].values()}
        exp |= {_norm(m) for m in re.findall(r"\{([^}]+)\}", tpl["out"])}
    return exp

def score_columns(cols) -> int:
    exp = expected_headers()
    return sum(1 for c in cols if _norm(c) in exp)

def read_wide_try(file_bytes: bytes, is_xlsx: bool, header_row: int) -> Tuple[pd.DataFrame, Dict]:
    if is_xlsx:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, header=max(header_row-1,0))
        return df, {"source":"xlsx", "mode":"wide", "header_row": header_row-1}
    else:
        sample = file_bytes[:2048].decode("utf-8", errors="ignore")
        try: sep = csv.Sniffer().sniff(sample).delimiter
        except Exception: sep = ","
        df = pd.read_csv(io.BytesIO(file_bytes), sep=sep, header=max(header_row-1,0))
        return df, {"source":"csv", "mode":"wide", "header_row": header_row-1}

def read_kv_from_raw(file_bytes: bytes, is_xlsx: bool, key_row: int = 1, val_row: int = 2) -> Tuple[Dict[str,str], Dict]:
    if is_xlsx:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, header=None)
    else:
        df = pd.read_csv(io.BytesIO(file_bytes), header=None)
    keys = [safe(x).replace("\ufeff","").replace("\xa0"," ") for x in df.iloc[key_row-1].tolist()]
    vals = [safe(x).replace("\ufeff","").replace("\xa0"," ") for x in df.iloc[val_row-1].tolist()]
    kv = {k: v for k, v in zip(keys, vals) if k}
    return kv, {"source":"xlsx" if is_xlsx else "csv", "mode":"kv", "key_row":key_row-1, "val_row":val_row-1}

def extract_record_from_upload(file: UploadFile, header_row: int) -> Tuple[Dict[str,str], Dict, Optional[list]]:
    data = file.file.read()
    name = (file.filename or "").lower()
    is_xlsx = name.endswith(".xlsx")
    if not (is_xlsx or name.endswith(".csv")):
        raise HTTPException(400, "Поддерживаются только .xlsx или .csv")

    df_wide, meta = read_wide_try(data, is_xlsx, header_row)
    if not df_wide.empty:
        cols = [str(c) for c in df_wide.columns]
        sc = score_columns(cols)  # теперь только для информации
        row = pick_first_nonempty_row(df_wide)
        row_dict = {str(k): safe(v) for k, v in row.items()}
        meta.update({"mode": "wide", "score": sc})
        return row_dict, meta, cols

    # если df_wide пустой (совсем ничего не прочитали) — пробуем kv-режим
    kv, meta_kv = read_kv_from_raw(data, is_xlsx, 1, 2)
    meta_kv.setdefault("score", 0)
    return kv, meta_kv, None

def extract_record_from_gsheet(url: str, header_row: int) -> Tuple[Dict[str,str], Dict, Optional[list]]:
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", url or "")
    if not m: raise HTTPException(400, "Не удалось извлечь spreadsheetId из URL")
    spreadsheet_id = m.group(1)
    gid_match = re.search(r"[#&?]gid=([0-9]+)", url)
    gid = int(gid_match.group(1)) if gid_match else 0
    export = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=csv&gid={gid}"
    resp = requests.get(export, timeout=30)
    if resp.status_code != 200:
        raise HTTPException(400, f"Google Sheets недоступен (HTTP {resp.status_code})")
    upl = UploadFile(filename="gs.csv", file=io.BytesIO(resp.content))
    rec, meta, cols = extract_record_from_upload(upl, header_row)
    meta.update({"source":"gsheet", "gid": gid})
    return rec, meta, cols

def pick_first_nonempty_row(df: pd.DataFrame) -> pd.Series:
    df = df.fillna("")
    for _, row in df.iterrows():
        if any(safe(v) for v in row.values):
            return row
    raise HTTPException(400, "Не найдена ни одна непустая строка с данными")

# -------- шаблон Excel --------
@app.get("/template")
def download_template(
    kit: Optional[str] = Query(
        default=None,
        description="id комплекта (kit1, kit2, kit3, kit4)",
    ),
    include: Optional[str] = Query(
        default=None,
        description="CSV-список id шаблонов (старый режим, можно не использовать)",
    ),
):
    kit = kit.strip()
    path = KIT_TEMPLATES.get(kit)
    if not path:
        raise HTTPException(
            400,
            detail=f"Неизвестный комплект: {kit}",
        )
    if not path.is_file():
        raise HTTPException(
            500,
            detail=f"Файл шаблона для комплекта {kit} не найден по пути {path}",
        )

    return FileResponse(
        path,
        filename=path.name,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

# -------- инструкция DOCX --------
def _build_instruction_docx_bytes() -> bytes:
    """Генерация дефолтной инструкции (если нет готового файла)."""
    doc = Document()
    # стиль
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Инструкция по заполнению Excel-таблицы main_example.xlsx", 0)

    p = doc.add_paragraph(
        "Таблица содержит один лист. В первой строке расположены названия полей, "
        "во второй строке — значения для одного студента. На основании этих значений "
        "формируются все документы Word из набора шаблонов."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Общие правила", level=1)
    rules = [
        "Формат даты: ДД.ММ.ГГГГ (например, 17.10.2025).",
        "Телефон в формате +7 999 123-45-67 или 8 999 123-45-67.",
        "E-mail: user@example.org.",
        "ФИО в именительном падеже: Иванов Иван Иванович.",
        "Поле «Курс» — целое число (1, 2, 3, 4...).",
        "ИНН организации — 10 или 12 цифр, без пробелов.",
        "Адреса указывайте полностью, как в официальных документах.",
        "Если колонка продублирована (например, «АдресОрганизации» и «АдрессОрганизации»), укажите одно и то же значение в обеих.",
    ]
    for r in rules:
        doc.add_paragraph(r, style=None).paragraph_format.left_indent = Pt(14)

    doc.add_heading("Список основных полей", level=1)
    fields = [
        ("ФИО", "Полное имя студента (И.П.). Пример: Иванов Иван Иванович."),
        ("Группа", "Учебная группа. Пример: Изу-101."),
        ("Курс", "Номер курса. Пример: 3."),
        ("ТипПрактики / ВидПрактика", "Например: производственная, преддипломная."),
        ("НачалоПрактики / КонецПрактики", "Даты в формате ДД.ММ.ГГГГ."),
        ("БазаПрактики", "Наименование организации. Пример: ООО «Ромашка»."),
        ("АдресОрганизации / АдрессОрганизации", "Почтовый адрес организации (одно и то же значение в обеих колонках)."),
        ("ЮрАдресПрофОрг", "Юридический адрес учебного подразделения."),
        ("ОргИНН", "ИНН организации."),
        ("РукПрофОрг / РукВУЗФИО / РукВУЗ", "ФИО/должности руководителей от организации и вуза."),
        ("Кафедра / КафедраРП", "Название кафедры."),
        ("Научный руководитель", "ФИО, должность, степень, звание: ФИОНаучРук, ДолжНаучРук, СтепеньНаучРук, ЗваниеНаучРук."),
        ("ФИОДП", "ФИО студента в требуемом падеже для ВКР (обычно родительный)."),
        ("СегодняшняяДата", "Текущая дата формирования документов."),
    ]
    for name, desc in fields:
        doc.add_paragraph(f"• {name}: {desc}")

    doc.add_heading("Имена выходных файлов", level=1)
    doc.add_paragraph(
        "Имена документов формируются автоматически и включают ФИО и группу, "
        "например: «Дневник_{ФИО}_{Группа}.docx», «Титул_ВКР_{ФИО}_{Группа}.docx»."
    )

    doc.add_heading("Где используются данные", level=1)
    doc.add_paragraph(
        "Полный перечень соответствий «поле → документ(ы)» указан на странице сервиса ниже в инструкции. "
        "Заполняйте все поля без пропусков — пустые ячейки приводят к незаполненным местам в результатах."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

@app.get("/instruction")
def download_instruction():
    """
    Отдаём инструкцию (DOCX). Если в корне лежит готовый файл (instruction.docx / инструкция.docx / …),
    вернём его. Иначе — сгенерируем типовой DOCX на лету.
    """
    for p in INSTRUCTION_CANDIDATES:
        if p.exists():
            return FileResponse(
                str(p),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=INSTRUCTION_DOWNLOAD_NAME,
                headers={"Cache-Control": "no-store, no-cache, must-revalidate"},
            )
    # fallback: сгенерируем docx
    content = _build_instruction_docx_bytes()
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{INSTRUCTION_DOWNLOAD_NAME}"',
            "Cache-Control": "no-store, no-cache, must-revalidate",
        },
    )

# ============= HTTP API =============
@app.get("/", response_class=HTMLResponse)
def index():
    return HTMLResponse(INDEX_HTML)

@app.post("/inspect")
def inspect(
    table_file: Optional[UploadFile] = File(default=None),
    gsheet_url: Optional[str] = Form(default=None),
    header_row: int = Form(default=1),
):
    # приоритет: если есть ссылка — используем её, иначе файл
    if gsheet_url and gsheet_url.strip():
        record, meta, cols = extract_record_from_gsheet(gsheet_url.strip(), header_row)
    elif table_file and (table_file.filename or "").strip():
        record, meta, cols = extract_record_from_upload(table_file, header_row)
    else:
        raise HTTPException(400, "Укажите Google Sheet ИЛИ выберите файл")

    needed = ["ФИО", "Группа"]
    missing = [k for k in needed if k not in record]

    if meta["mode"] == "wide":
        preview = record
        return JSONResponse({"columns": cols or [], "preview": preview, "missing": missing, "meta": meta})
    else:
        preview_pairs = list(record.items())[:12]
        return JSONResponse({"columns": [], "preview_pairs": preview_pairs, "missing": missing, "meta": meta})

@app.post("/generate")
def generate_zip(
    table_file: Optional[UploadFile] = File(default=None),
    gsheet_url: Optional[str] = Form(default=None),
    header_row: int = Form(default=1),
    include: Optional[str] = Form(default=None), 
):
    if gsheet_url and gsheet_url.strip():
        record, meta, _ = extract_record_from_gsheet(gsheet_url.strip(), header_row)
    elif table_file and (table_file.filename or "").strip():
        record, meta, _ = extract_record_from_upload(table_file, header_row)
    else:
        raise HTTPException(400, "Укажите Google Sheet ИЛИ выберите файл")

    fio = safe(record.get("ФИО")) or "record"
    folder = slugify(f"001_{fio}")

    selected_ids = None
    if include:
        selected_ids = {
            s.strip().lower()
            for s in include.split(",")
            if s.strip()
        }

    templates = TEMPLATES
    if selected_ids:
        templates = [t for t in TEMPLATES if t.get("id") in selected_ids]
        # если вдруг кто-то к нам постучался с левыми id — просто сгенерим пустой ZIP с ошибками
        if not templates:
            return JSONResponse(
                {
                    "error": "Ни один шаблон не совпал с include",
                    "include": sorted(selected_ids),
                    "available": [t["id"] for t in TEMPLATES],
                },
                status_code=400,
            )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for tpl in templates:
            try:
                # контекст: {tpl_key: значение из record по названию колонки}
                ctx = {tpl_key: safe(record.get(excel_col, "")) for tpl_key, excel_col in tpl["fields"].items()}
                doc = DocxTemplate(tpl["path"])
                doc.render(ctx, jinja_env=JINJA_ENV)

                # рендерим DOCX в память
                out_mem = io.BytesIO()
                doc.save(out_mem)
                docx_bytes = out_mem.getvalue()

                # имя файла из шаблонной маски out
                out_name = slugify(tpl["out"].format_map(SafeMap(record)) or "doc_001.docx")

                # формат выхода: по умолчанию docx, но для выбранных шаблонов = pdf
                output = (tpl.get("output") or "docx").strip().lower()
                if output == "pdf":
                    # гарантируем расширение .pdf
                    if out_name.lower().endswith(".docx"):
                        out_name = out_name[:-5] + ".pdf"
                    elif not out_name.lower().endswith(".pdf"):
                        out_name += ".pdf"

                # собираем путь внутри архива (с учётом подпапки dir)
                subdir_raw = (tpl.get("dir") or "").strip()
                if subdir_raw:
                    subdir_filled = slugify_path(subdir_raw.format_map(SafeMap(record)))
                    arcname = "/".join([folder, subdir_filled, out_name])
                else:
                    arcname = "/".join([folder, out_name])

                # пишем либо pdf, либо docx
                if output == "pdf":
                    pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)
                    zf.writestr(arcname, pdf_bytes)
                else:
                    # гарантируем расширение .docx
                    if not out_name.lower().endswith(".docx"):
                        # (на случай если в конфиге забыли расширение)
                        arcname = arcname + ".docx"
                    zf.writestr(arcname, docx_bytes)
            except Exception as e:
                err = slugify(tpl.get("out","file")) + ".ERROR.txt"
                zf.writestr(f"{folder}/{err}", f"Ошибка ({tpl['path']}): {type(e).__name__}: {e}")

    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="generated_docs.zip"'}
    )

@app.get("/healthz")
def healthz():
    return PlainTextResponse("ok")
